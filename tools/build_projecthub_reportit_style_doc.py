from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_projecthub_api_doc import BASE_URL, build_entries


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "ProjectHub_Neat_ReportIt_Style_Documentation.docx"


def set_run_font(run, name="Times New Roman", size=12):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)


def para(doc, text="", bold=False, size=12, align=None, before=0, after=2, left=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if left:
        p.paragraph_format.left_indent = Inches(left)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size)
    r.bold = bold
    return p


def label(doc, text):
    return para(doc, text, bold=True, before=3, after=0)


def code_block(doc, value):
    if value is None:
        text = ""
    elif isinstance(value, str):
        text = value
    else:
        import json

        text = json.dumps(value, indent=2, ensure_ascii=False)
    if text == "":
        return para(doc, "", after=0)
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.22)
        r = p.add_run(line)
        set_run_font(r, "Courier New", 10)


def business_logic(entry):
    if entry.get("logic"):
        return entry["logic"]
    purpose = entry["purpose"].rstrip(".")
    return [
        f"Validate request data for {entry['path']}.",
        f"Perform service-layer operation to {purpose.lower()}.",
        "Return success data or an error response.",
    ]


def add_front_matter(doc, count):
    para(doc, "PROJECTHUB API DOCUMENTATION", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(
        doc,
        "Student Project Submission and Review Platform",
        bold=False,
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
    )
    para(doc)

    para(doc, "1. INTRODUCTION", bold=True, before=8, after=4)
    para(
        doc,
        "ProjectHub is a student project submission, certificate verification, faculty review, "
        "and admin management platform. Students can upload projects and certificates, faculty "
        "members can review assigned work, and admins can manage users, settings, assignments, "
        "deadlines, analytics, and audit data.",
    )
    para(doc, "The system follows a Spring Boot REST API backend with:")
    para(doc, "MySQL for users, student profiles, faculty profiles, settings, and audit logs.")
    para(doc, "MongoDB for project, certificate, notification, and deadline-rule documents.")
    para(doc, "JWT authentication for secured backend APIs.")
    para(doc, f"This document includes {count} ProjectHub APIs.")
    para(doc)

    para(doc, "2. AUTHENTICATION STANDARDS", bold=True, before=8, after=4)
    label(doc, "Authentication Method")
    para(doc, "Secured endpoints require JWT authentication.")
    label(doc, "Header")
    para(doc, "Authorization: Bearer <token>")
    label(doc, "Example")
    para(doc, "Authorization: Bearer eyJhbGciOiJIUzI1NiJ9...")
    label(doc, "Base URL")
    para(doc, BASE_URL)
    label(doc, "Service Mode Note")
    para(
        doc,
        "The backend has auth mode and backend mode. Login, register, forgot-password, and "
        "reset-password are public in auth mode. Normal project, certificate, admin, and profile "
        "APIs require JWT/roles in backend mode.",
    )
    para(doc)

    para(doc, "3. COMMON RESPONSE FORMATS", bold=True, before=8, after=4)
    label(doc, "Success Response")
    code_block(doc, {"message": "Operation completed successfully"})
    label(doc, "Error Response")
    code_block(doc, {"error": "Invalid request"})
    label(doc, "Validation Error")
    code_block(doc, {"error": "Validation failed", "details": {"field": "Field is required"}})
    para(doc)


def add_api_entry(doc, api_no, entry):
    title = entry["purpose"].rstrip(".").upper()
    para(doc, f"API {api_no}: {title}", bold=True, before=8, after=3)
    label(doc, "Endpoint")
    para(doc, f"{entry['method']} {entry['path']}")
    label(doc, "Full URL")
    para(doc, entry.get("url") or f"{BASE_URL}{entry['path']}")
    label(doc, "Description")
    para(doc, entry["purpose"])
    label(doc, "Authentication")
    para(doc, entry["auth"])
    label(doc, "Request Headers")
    auth_text = entry["auth"].lower()
    is_public = "public" in auth_text or "no jwt" in auth_text
    if is_public:
        if "multipart" in (entry.get("notes") or "").lower() or "upload" in entry["path"]:
            para(doc, "Content-Type: multipart/form-data")
        elif entry["method"] != "GET":
            para(doc, "Content-Type: application/json")
        else:
            para(doc, "No authentication header required")
    elif entry["requestJson"] is None and entry["method"] == "GET":
        para(doc, "Authorization: Bearer <token> for secured endpoints")
    elif "multipart" in (entry.get("notes") or "").lower() or "upload" in entry["path"]:
        para(doc, "Content-Type: multipart/form-data")
        para(doc, "Authorization: Bearer <token> for secured endpoints")
    else:
        para(doc, "Content-Type: application/json")
        para(doc, "Authorization: Bearer <token> for secured endpoints")
    label(doc, "Request Body")
    code_block(doc, entry["requestJson"])
    label(doc, "Request Parameters")
    code_block(doc, entry.get("queryOrPathParams") or "")
    label(doc, "Success Response")
    code_block(doc, entry["responseJson"])
    label(doc, "Status Codes")
    para(doc, "200 OK - Request completed successfully")
    para(doc, "400 BAD REQUEST - Invalid input or missing required data")
    para(doc, "401 UNAUTHORIZED - Missing or invalid JWT token")
    para(doc, "403 FORBIDDEN - User role is not allowed")
    para(doc, "500 INTERNAL SERVER ERROR - Server-side failure")
    label(doc, "Business Logic")
    for line in business_logic(entry):
        para(doc, line, left=0.18)
    para(doc, after=4)


def module_title(module_no, module):
    return f"MODULE {module_no} - {module.upper()} APIs"


def build_doc():
    entries, _ = build_entries()
    grouped = {}
    for entry in entries:
        grouped.setdefault(entry["module"], []).append(entry)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.08

    add_front_matter(doc, len(entries))

    api_no = 1
    module_no = 1
    for module, module_entries in grouped.items():
        if module_no > 1:
            doc.add_page_break()
        para(doc, module_title(module_no, module), bold=True, before=6, after=6)
        para(doc)
        for entry in module_entries:
            add_api_entry(doc, api_no, entry)
            api_no += 1
        module_no += 1

    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX}")
    print(f"apis {len(entries)}")


if __name__ == "__main__":
    build_doc()
