import json
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_JSON = ROOT / "projecthub--1-main" / "docs" / "ProjectHub_API_ModuleWise.json"
OUT_DOCX = ROOT / "ProjectHub_Complete_API_Documentation.docx"
BASE_URL = "http://localhost:8081"


def load_existing_entries():
    data = json.loads(SOURCE_JSON.read_text(encoding="utf-8"))
    entries = []
    for module, module_entries in data["modules"].items():
        entries.extend(module_entries)
    return entries, data


def endpoint_key(entry):
    return (entry["method"].upper(), entry["path"].split("?")[0])


def normalize_entry(entry):
    entry = dict(entry)
    entry["method"] = entry["method"].upper()
    entry.setdefault("url", BASE_URL + entry["path"])
    entry.setdefault("auth", "Bearer JWT required after login")
    entry.setdefault("queryOrPathParams", {})
    entry.setdefault("requestJson", None)
    entry.setdefault("responseJson", None)
    entry.setdefault("notes", "")
    return entry


SUPPLEMENTAL = [
    {
        "module": "Authentication",
        "method": "POST",
        "path": "/users/forgot-password",
        "auth": "Public in auth service mode; denied in backend service mode",
        "purpose": "Request a student password reset token for the submitted email.",
        "requestJson": {"email": "student@test.com"},
        "responseJson": {"success": True, "message": "Password reset token generated or sent"},
        "notes": "Implemented by UserController and secured differently depending on app.service-mode.",
        "logic": ["Read email from request body.", "Create or send a password reset token for the student account.", "Return a map with reset status details."],
    },
    {
        "module": "Authentication",
        "method": "POST",
        "path": "/users/reset-password",
        "auth": "Public in auth service mode; denied in backend service mode",
        "purpose": "Reset a student password using a reset token.",
        "requestJson": {"token": "reset-token", "newPassword": "NewPassword@123", "confirmPassword": "NewPassword@123"},
        "responseJson": {"success": True, "message": "Password reset successful"},
        "logic": ["Validate reset token.", "Check new password and confirmation match.", "Update the student password."],
    },
    {
        "module": "Faculty Assignment",
        "method": "GET",
        "path": "/users/faculty/{facultyUserId}/students",
        "auth": "Bearer JWT required after login",
        "purpose": "Get students assigned to a faculty user.",
        "queryOrPathParams": {"facultyUserId": 38},
        "requestJson": None,
        "responseJson": [{"id": 2, "department": "CSE", "studentCode": "STU-S157", "faculty": {"id": 2}}],
        "logic": ["Resolve the faculty record from faculty user id.", "Return all student profiles mapped to that faculty."],
    },
    {
        "module": "Faculty Assignment",
        "method": "PUT",
        "path": "/users/admin/assign-student?studentUserId=30&facultyUserId=38",
        "auth": "ADMIN role required",
        "purpose": "Assign one student to one faculty member.",
        "queryOrPathParams": {"studentUserId": 30, "facultyUserId": 38},
        "requestJson": None,
        "responseJson": {"id": 2, "studentCode": "STU-S157", "faculty": {"id": 2, "facultyCode": "FAC-001"}},
        "logic": ["Find the student profile by student user id.", "Find the faculty profile by faculty user id.", "Persist the student-faculty assignment."],
    },
    {
        "module": "Faculty Assignment",
        "method": "PUT",
        "path": "/users/admin/unassign-student?studentUserId=30",
        "auth": "ADMIN role required",
        "purpose": "Remove a student's faculty assignment.",
        "queryOrPathParams": {"studentUserId": 30},
        "requestJson": None,
        "responseJson": {"id": 2, "studentCode": "STU-S157", "faculty": None},
        "logic": ["Find the student profile by user id.", "Clear the faculty relation.", "Save the updated student profile."],
    },
    {
        "module": "Faculty Assignment",
        "method": "PUT",
        "path": "/users/admin/assign-students-by-department?department=CSE&section=A&facultyUserId=38",
        "auth": "ADMIN role required",
        "purpose": "Assign all students in a department, optionally filtered by section, to a faculty member.",
        "queryOrPathParams": {"department": "CSE", "section": "A", "facultyUserId": 38},
        "requestJson": None,
        "responseJson": [{"id": 2, "department": "CSE", "section": "A", "faculty": {"id": 2}}],
        "logic": ["Filter student profiles by department and optional section.", "Resolve faculty by user id.", "Assign matching students to the faculty profile."],
    },
    {
        "module": "Faculty Assignment",
        "method": "GET",
        "path": "/users/students/department/{department}",
        "auth": "ADMIN role required",
        "purpose": "Get student profiles for a department.",
        "queryOrPathParams": {"department": "CSE"},
        "requestJson": None,
        "responseJson": [{"id": 2, "department": "CSE", "studentCode": "STU-S157"}],
        "logic": ["Read department from the path.", "Return all matching student profiles."],
    },
    {
        "module": "Landing",
        "method": "GET",
        "path": "/landing/click?action=hero_cta",
        "auth": "Public / no JWT required in backend service mode",
        "purpose": "Track a landing-page click action and return the tracked action with timestamp.",
        "queryOrPathParams": {"action": "hero_cta"},
        "requestJson": None,
        "responseJson": {"status": "tracked", "action": "hero_cta", "time": "2026-06-15T10:30:00"},
        "logic": ["Read the action query parameter, defaulting to unknown.", "Return a simple tracking response with server time."],
    },
]


MODULE_ORDER = [
    "Authentication",
    "Users",
    "Profiles",
    "Faculty Assignment",
    "Settings and Analytics",
    "Notifications",
    "Audit Logs",
    "Projects",
    "Certificates",
    "Deadline Rules",
    "Admin",
    "Dashboard",
    "Landing",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def add_code_paragraph(doc, value):
    if value is None:
        text = "None"
    elif isinstance(value, str):
        text = value
    else:
        text = json.dumps(value, indent=2, ensure_ascii=False)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for line_no, line in enumerate(text.splitlines() or [""]):
        if line_no:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(31, 45, 61)


def add_label_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_cell_margins(table)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], "E8EEF5")
        for p in cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
                set_run_font(run)
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run)
                    run.font.size = Pt(9.5)
    set_table_width(table, [1.55, 4.95])
    doc.add_paragraph()


def add_endpoint_index(doc, entries):
    doc.add_heading("API URL Index", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_cell_margins(table)
    headers = ["No.", "Module", "Method", "Endpoint", "Auth"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
    for idx, entry in enumerate(entries, 1):
        row = table.add_row().cells
        values = [str(idx), entry["module"], entry["method"], entry["path"], entry["auth"]]
        for cell, value in zip(row, values):
            cell.text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run)
                    run.font.size = Pt(8)
                    if row is table.rows[0]:
                        run.bold = True
    set_table_width(table, [0.42, 1.15, 0.62, 2.65, 1.66])


def add_standard_sections(doc, total_count):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROJECTHUB API DOCUMENTATION")
    set_run_font(run)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Student project submission, review, certificate, assignment, and admin platform")
    set_run_font(run)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(85, 85, 85)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Base URL: {BASE_URL} | Total APIs documented: {total_count}")
    set_run_font(run)
    run.font.size = Pt(10)

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "ProjectHub is a student project management platform for submitting projects, "
        "reviewing work, verifying certificates, managing faculty assignments, and tracking "
        "administrative analytics. The backend is a Spring Boot REST API using MySQL for user "
        "and profile data, MongoDB for project/certificate/notification records, and JWT-based "
        "security for protected endpoints."
    )
    doc.add_paragraph("The system modules covered in this document are:")
    for item in [
        "Authentication and password reset",
        "Users, student profiles, faculty profiles, and assignment workflows",
        "Project upload, MongoDB project CRUD, search, review, and student analytics",
        "Certificate submission, verification, resubmission, and deletion",
        "Admin settings, dashboard analytics, notifications, audit logs, and deadline rules",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Authentication Standards", level=1)
    add_label_value_table(
        doc,
        [
            ("Authentication method", "JWT bearer token for secured endpoints."),
            ("Header", "Authorization: Bearer <token>"),
            ("Content type", "application/json unless the endpoint explicitly uses multipart/form-data."),
            ("Service modes", "auth mode exposes login/register/reset APIs; backend mode protects normal business APIs and permits public downloads/landing tracking."),
        ],
    )

    doc.add_heading("3. Common Response Formats", level=1)
    doc.add_paragraph("Success response example")
    add_code_paragraph(doc, {"message": "Operation completed successfully"})
    doc.add_paragraph("Error response example")
    add_code_paragraph(doc, {"error": "Invalid request"})
    doc.add_paragraph("Validation error example")
    add_code_paragraph(doc, {"error": "Validation failed", "details": {"field": "Field is required"}})


def apply_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def sort_entries(entries):
    order = {name: idx for idx, name in enumerate(MODULE_ORDER)}
    return sorted(entries, key=lambda e: (order.get(e["module"], 99), e["module"], e["method"], e["path"]))


def build_entries():
    existing, source_data = load_existing_entries()
    entries_by_key = OrderedDict()
    for entry in existing:
        normalized = normalize_entry(entry)
        entries_by_key[endpoint_key(normalized)] = normalized
    for entry in SUPPLEMENTAL:
        normalized = normalize_entry(entry)
        entries_by_key[endpoint_key(normalized)] = normalized
    return sort_entries(list(entries_by_key.values())), source_data


def add_endpoint_details(doc, entries):
    grouped = OrderedDict()
    for entry in entries:
        grouped.setdefault(entry["module"], []).append(entry)

    doc.add_heading("4. Module-wise API Details", level=1)
    api_no = 1
    for module, module_entries in grouped.items():
        doc.add_heading(f"Module: {module}", level=2)
        for entry in module_entries:
            path_for_title = entry["path"]
            doc.add_heading(f"API {api_no}: {entry['method']} {path_for_title}", level=3)
            add_label_value_table(
                doc,
                [
                    ("Endpoint", f"{entry['method']} {entry['path']}"),
                    ("Full URL", entry["url"] if "url" in entry else BASE_URL + entry["path"]),
                    ("Description", entry["purpose"]),
                    ("Authentication", entry["auth"]),
                    ("Parameters", json.dumps(entry["queryOrPathParams"], ensure_ascii=False) if entry["queryOrPathParams"] else "None"),
                    ("Notes", entry.get("notes") or "None"),
                ],
            )
            doc.add_paragraph("Request Body / Form Data")
            add_code_paragraph(doc, entry["requestJson"])
            doc.add_paragraph("Success Response")
            add_code_paragraph(doc, entry["responseJson"])
            doc.add_paragraph("Business Logic")
            logic = entry.get("logic")
            if not logic:
                logic = [
                    f"Validate request data for {entry['path']}.",
                    "Call the matching service-layer method.",
                    "Return entity data, list data, status text, or analytics DTO as applicable.",
                ]
            for line in logic:
                doc.add_paragraph(line, style="List Bullet")
            api_no += 1


def add_appendix(doc, source_data):
    doc.add_heading("5. Data Samples", level=1)
    doc.add_paragraph("Representative MySQL sample")
    add_code_paragraph(doc, source_data.get("sqlSamples", {}).get("student_user"))
    doc.add_paragraph("Representative MongoDB project sample")
    add_code_paragraph(doc, source_data.get("mongoSamples", {}).get("project"))
    doc.add_paragraph("Representative certificate sample")
    add_code_paragraph(doc, source_data.get("mongoSamples", {}).get("certificate"))


def main():
    entries, source_data = build_entries()
    doc = Document()
    apply_styles(doc)
    add_standard_sections(doc, len(entries))
    add_endpoint_index(doc, entries)
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_endpoint_details(doc, entries)
    add_appendix(doc, source_data)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("ProjectHub API Documentation")
    set_run_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX}")
    print(f"apis {len(entries)}")


if __name__ == "__main__":
    main()
