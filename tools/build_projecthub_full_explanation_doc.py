import json
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_projecthub_api_doc import build_entries


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "ProjectHub_Annotations_JWT_API_Buttons_Explanation.docx"


def set_font(run, name="Calibri", size=11, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph(doc, text="", bold=False, size=11, color="000000", align=None, before=0, after=5, left=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if left:
        p.paragraph_format.left_indent = Inches(left)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, color=color)
    r.bold = bold
    return p


def heading(doc, text, level=1):
    sizes = {1: 16, 2: 13, 3: 11}
    p = paragraph(doc, text, bold=True, size=sizes.get(level, 11), color="000000", before=10 if level == 1 else 6, after=5)
    return p


def bullet(doc, text, left=0.22):
    paragraph(doc, text, size=10.5, after=3, left=left)


def code(doc, value):
    if value is None:
        text = "No request body"
    elif isinstance(value, str):
        text = value
    else:
        text = json.dumps(value, indent=2, ensure_ascii=False)
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_font(r, "Consolas", 9)


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_font(r, size=size)
    r.bold = bold


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        shade(t.rows[0].cells[i], "F2F2F2")
        set_cell_text(t.rows[0].cells[i], h, bold=True, size=9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=8.5)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    paragraph(doc, "", after=2)
    return t


ANNOTATIONS = [
    ("@SpringBootApplication", "Main app annotation", "Starts Spring Boot auto-configuration, component scanning, and application bootstrapping."),
    ("@RestController", "Controller layer", "Marks a class as a REST API controller. Methods return JSON/text responses directly."),
    ("@RequestMapping", "Base route", "Sets the common URL prefix for all APIs in that controller, for example /users or /projects."),
    ("@GetMapping", "Read API", "Maps an HTTP GET request. Used for fetching users, profiles, projects, certificates, settings, analytics, and downloads."),
    ("@PostMapping", "Create/login/upload API", "Maps an HTTP POST request. Used for login, registration, create records, upload files, and save settings."),
    ("@PutMapping", "Update/action API", "Maps an HTTP PUT request. Used for profile update, review approval/rejection, resubmit, mark notification read, and assignments."),
    ("@DeleteMapping", "Delete API", "Maps an HTTP DELETE request. Used for deleting users, projects, certificates, notifications, and deadline rules."),
    ("@CrossOrigin(\"*\")", "CORS", "Allows the React frontend to call the backend APIs from another port/domain."),
    ("@Autowired", "Dependency injection", "Injects service, repository, filter, or utility objects automatically into a class."),
    ("@Service", "Service layer", "Marks a class that contains business logic, such as UserService and ProjectService."),
    ("@Repository", "Data layer", "Marks database access interfaces/classes. Spring creates implementations for JPA/Mongo repositories."),
    ("@Entity", "MySQL entity", "Maps a Java class to a relational database table, for example User, Student, Faculty, AdminSettings."),
    ("@Document", "MongoDB document", "Maps a Java class to a MongoDB collection, for example projects, certificates, notifications, deadline_rules."),
    ("@Id", "Primary key", "Marks the unique id field of a MySQL entity or MongoDB document."),
    ("@GeneratedValue", "Auto id generation", "Lets the database generate MySQL primary key values automatically."),
    ("@OneToOne", "One-to-one relation", "Defines one linked record, for example one User has one Student or Faculty profile."),
    ("@ManyToOne", "Many-to-one relation", "Defines many records linked to one record, for example many Students assigned to one Faculty."),
    ("@JoinColumn", "Foreign key column", "Names the database column used to join related MySQL tables."),
    ("@RequestBody", "JSON body input", "Reads JSON data from the frontend request body into an object or map."),
    ("@RequestParam", "Query/form input", "Reads query parameters or multipart form fields, for example status, feedback, grade, file."),
    ("@PathVariable", "URL path input", "Reads dynamic values from the URL, for example /users/{id} or /projects/mongo/{id}."),
    ("@Valid", "Validation trigger", "Runs validation rules on a request body before the controller method continues."),
    ("@Configuration", "Configuration class", "Marks classes that define Spring configuration, like SecurityConfig and PasswordConfig."),
    ("@Bean", "Spring bean method", "Creates objects managed by Spring, like SecurityFilterChain and PasswordEncoder."),
    ("@Value", "Property injection", "Reads values from application properties/environment, such as app.service-mode."),
    ("extends OncePerRequestFilter", "JWT filter base", "Ensures JWTFilter runs once per HTTP request to validate tokens."),
]


JWT_STEPS = [
    ("1. Login request", "User enters email and password in Login.jsx. Frontend sends POST /users/login with AuthRequest data."),
    ("2. Credential validation", "UserController.login() calls UserService.login(), which checks the user record and password."),
    ("3. Token generation", "JWTUtil.generateToken(email, role) creates a JWT with subject=email, role claim, issuedAt, expiry, and HS256 signature."),
    ("4. Login response", "Backend returns AuthResponse containing token, id, role, and fullName."),
    ("5. Frontend storage", "React stores token, id, role, and fullName in localStorage for later API calls."),
    ("6. Protected API call", "Frontend sends Authorization: Bearer <token> in headers for protected APIs."),
    ("7. JWTFilter check", "JWTFilter runs before controller. Public URLs are skipped; protected URLs require a valid bearer token."),
    ("8. Token validation", "JWTUtil validates signature and expiry, extracts email and role claim."),
    ("9. Security context", "JWTFilter creates UsernamePasswordAuthenticationToken with ROLE_STUDENT, ROLE_FACULTY, or ROLE_ADMIN."),
    ("10. Controller execution", "If authentication succeeds, the request reaches the controller. If missing/invalid/expired, backend returns 401."),
    ("11. Logout", "Logout clears localStorage and redirects to Login. Since JWT is stateless, no server session is destroyed."),
]


BUTTON_ACTIONS = OrderedDict({
    "Login / Home": [
        ("Login button", "POST /users/login", "Authenticates user and stores JWT details in localStorage."),
        ("Register / Signup action", "POST /auth/register or POST /users", "Creates a new account depending on service mode."),
        ("Forgot Password", "POST /users/forgot-password", "Requests password reset for a student email."),
        ("Reset Password", "POST /users/reset-password", "Sets a new password using reset token."),
        ("Landing click / CTA", "GET /landing/click?action=...", "Tracks landing page button click."),
    ],
    "Student Dashboard": [
        ("Dashboard refresh", "GET /dashboard and GET /projects/analytics/student/{studentId}", "Loads student counts and analytics."),
        ("Submit Project navigation", "No backend API directly", "Opens Submit Project page."),
        ("Certificates navigation", "No backend API directly", "Opens certificate upload/list page."),
        ("Profile navigation", "GET /users/student/profile/{id}", "Loads student profile screen."),
        ("Notifications navigation", "GET /users/notifications/{userId}", "Loads notifications."),
        ("Logout", "Client-side localStorage clear", "Removes token and redirects to Login."),
    ],
    "Submit Project": [
        ("New Submission", "No backend API until Submit", "Opens project form."),
        ("Submit Project", "POST /projects/upload then POST /projects/mongo", "Uploads file and creates MongoDB project record."),
        ("Edit Project", "PUT /projects/mongo/update/{id}", "Updates existing project details."),
        ("Delete Project", "DELETE /projects/mongo/{id}", "Deletes selected project."),
        ("Resubmit", "PUT /projects/mongo/resubmit/{id}", "Increments version and returns project to pending workflow."),
        ("View / Download File", "GET /projects/download/{fileName}", "Opens or downloads submitted project file."),
        ("Download Feedback", "Client-side generated text/PDF", "Downloads faculty feedback details from project record."),
        ("Notify Faculty", "POST /users/notifications", "Creates a notification after submission/update."),
    ],
    "Certificates": [
        ("Upload Certificate", "POST /projects/upload then POST /certificates", "Uploads certificate file and creates certificate record."),
        ("Edit Certificate", "PUT /certificates/{id}", "Updates certificate details."),
        ("Delete Certificate", "DELETE /certificates/{id}", "Deletes selected certificate."),
        ("View / Download Certificate", "GET /projects/download/{fileName}", "Opens certificate file."),
        ("Notify Faculty", "POST /users/notifications", "Creates notification for assigned faculty."),
    ],
    "Profile": [
        ("Load Student Profile", "GET /users/student/profile/{id}", "Displays logged-in student profile."),
        ("Save Student Profile", "PUT /users/student/profile/{id}", "Updates department, year, links, and profile fields."),
        ("Load Faculty Profile", "GET /users/faculty/profile/{id}", "Displays faculty profile."),
        ("Save Faculty Profile", "PUT /users/faculty/profile/{id}", "Updates faculty department, specialization, section, and related fields."),
    ],
    "Notifications": [
        ("Open Notifications", "GET /users/notifications/{userId}", "Fetches notifications for logged-in user."),
        ("Mark Read", "PUT /users/notifications/read/{id}", "Marks one notification as read."),
        ("Mark All Read", "PUT /users/notifications/read/{id} for each unread item", "Marks all visible notifications as read."),
        ("Delete Notification", "DELETE /users/notifications/{id}", "Deletes notification from MongoDB."),
    ],
    "Faculty Dashboard / Review": [
        ("Faculty Dashboard load", "GET /projects/mongo/faculty/{facultyId}, GET /certificates/faculty/{facultyId}", "Loads faculty review summary."),
        ("Review Projects", "GET /projects/mongo/faculty/{facultyId}", "Lists projects assigned to faculty."),
        ("Approve Project", "PUT /projects/mongo/review/{id}?status=APPROVED&feedback=...&grade=...", "Approves project with feedback and grade."),
        ("Reject Project", "PUT /projects/mongo/review/{id}?status=REJECTED&feedback=...&grade=...", "Rejects project with feedback and grade."),
        ("Download Project File", "GET /projects/download/{fileName}", "Downloads project file for review."),
        ("Verify Certificates", "GET /certificates/faculty/{facultyId}", "Lists certificates assigned to faculty."),
        ("Approve Certificate", "PUT /certificates/{id}/verify?status=APPROVED&remarks=...", "Approves certificate."),
        ("Reject Certificate", "PUT /certificates/{id}/verify?status=REJECTED&remarks=...", "Rejects certificate."),
        ("Assigned Students", "GET /users/faculty/{facultyUserId}/students", "Lists students assigned to faculty."),
    ],
    "Admin Dashboard / Management": [
        ("Admin Dashboard load", "GET /users/admin/analytics and GET /users/audit-logs", "Loads admin analytics and recent activity."),
        ("Manage Users load", "GET /users", "Lists all users."),
        ("Create User", "POST /users", "Creates student, faculty, or admin user."),
        ("Update User", "PUT /users/{id}", "Updates user and profile fields."),
        ("Delete User", "DELETE /users/{id} or DELETE /admin/users/{id}", "Deletes selected user."),
        ("Search Users", "GET /users/search?name=...", "Searches users by name."),
        ("Filter Role", "GET /users/filter/role?role=...", "Filters users by role."),
        ("Bulk Upload", "POST /users/bulk-upload", "Creates users from Excel file."),
        ("View Submissions", "GET /admin/projects and GET /admin/certificates", "Shows all project and certificate submissions."),
        ("Admin Settings Save", "POST /users/settings", "Saves max file size, allowed file types, deadline settings."),
        ("Save Deadline", "POST /deadline-rules", "Creates or updates project/certificate deadline rules."),
        ("Assign Faculty Subject", "PUT /users/assign-faculty/{id}?subject=...", "Assigns subject to faculty user."),
        ("Assign Student to Faculty", "PUT /users/admin/assign-student?studentUserId=...&facultyUserId=...", "Assigns one student to faculty."),
        ("Assign by Department", "PUT /users/admin/assign-students-by-department?department=...&section=...&facultyUserId=...", "Assigns matching department/section students to faculty."),
    ],
})


def add_title(doc):
    paragraph(doc, "PROJECTHUB COMPLETE EXPLANATION DOCUMENTATION", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    paragraph(doc, "Annotations Used, JWT Workflow, Module APIs, and Frontend Buttons", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    paragraph(doc, "Prepared in simple English for easy project explanation and viva understanding.", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)


def add_annotations(doc):
    heading(doc, "1. Annotations Used in ProjectHub", 1)
    paragraph(doc, "This section explains the important Spring Boot, JPA, MongoDB, validation, and security annotations used in ProjectHub.", after=6)
    table(doc, ["Annotation", "Used For", "Meaning in ProjectHub"], ANNOTATIONS, widths=[1.7, 1.55, 3.25])


def add_jwt(doc):
    heading(doc, "2. JWT Workflow Used in ProjectHub", 1)
    paragraph(doc, "This section is based on the JWT workflow PDF you provided and the ProjectHub JWTUtil, JWTFilter, and SecurityConfig code.", after=6)
    heading(doc, "2.1 JWT Meaning", 2)
    paragraph(doc, "JWT means JSON Web Token. In ProjectHub, it is used to prove that a user has already logged in. The backend sends a signed token after login. The frontend sends that token again for protected APIs.")
    heading(doc, "2.2 Token Contents", 2)
    for item in [
        "Subject (sub): user's email.",
        "Custom claim role: STUDENT, FACULTY, or ADMIN.",
        "issuedAt: token creation time.",
        "expiration: current time plus 1 hour.",
        "Signature: HMAC-SHA256 using the secret key in JWTUtil.java.",
    ]:
        bullet(doc, item)
    heading(doc, "2.3 Step-by-step JWT Flow", 2)
    table(doc, ["Step", "Explanation"], JWT_STEPS, widths=[1.5, 5.0])
    heading(doc, "2.4 Public Endpoints That Skip JWTFilter", 2)
    for item in [
        "GET /projects/download/** - file downloads",
        "GET /landing/click - landing page tracking",
        "OPTIONS /** - CORS preflight",
        "POST /users/login, /users/forgot-password, /users/reset-password, /auth/login, /auth/register, /users",
        "/swagger-ui/**, /v3/api-docs/**, /swagger-ui.html, /error",
    ]:
        bullet(doc, item)
    heading(doc, "2.5 Protected Request Header", 2)
    code(doc, "Authorization: Bearer <token>")
    heading(doc, "2.6 Error Cases", 2)
    for item in [
        "If Authorization header is missing or not starting with Bearer, backend returns 401 JWT Token Required.",
        "If token is expired, tampered, or signed with the wrong key, backend returns 401 Invalid JWT Token.",
        "If token is valid but role is not allowed for the endpoint, Spring Security returns 403 Forbidden.",
    ]:
        bullet(doc, item)


def add_api_summary(doc, entries):
    heading(doc, "3. API List for All Modules", 1)
    paragraph(doc, "The following API list covers all ProjectHub modules from backend controllers and the generated API data.", after=6)
    grouped = OrderedDict()
    for entry in entries:
        grouped.setdefault(entry["module"], []).append(entry)
    api_no = 1
    for module, items in grouped.items():
        heading(doc, f"3.{list(grouped.keys()).index(module)+1} {module} APIs", 2)
        rows = []
        for entry in items:
            rows.append([api_no, entry["method"], entry["path"], entry["purpose"], entry["auth"]])
            api_no += 1
        table(doc, ["No.", "Method", "Endpoint", "Meaning / Purpose", "Auth"], rows, widths=[0.45, 0.65, 2.0, 2.1, 1.3])


def add_api_details(doc, entries):
    heading(doc, "4. Important API Details with Request and Response Meaning", 1)
    paragraph(doc, "This section gives understandable details for every API. Request and response examples are kept short where possible.", after=6)
    for idx, entry in enumerate(entries, 1):
        heading(doc, f"API {idx}: {entry['method']} {entry['path']}", 3)
        bullet(doc, f"Module: {entry['module']}")
        bullet(doc, f"Meaning: {entry['purpose']}")
        bullet(doc, f"Authentication: {entry['auth']}")
        if entry.get("queryOrPathParams"):
            bullet(doc, f"Parameters: {json.dumps(entry['queryOrPathParams'], ensure_ascii=False)}")
        paragraph(doc, "Request Example:", bold=True, after=1)
        code(doc, entry.get("requestJson"))
        paragraph(doc, "Response Example:", bold=True, after=1)
        code(doc, entry.get("responseJson"))


def add_buttons(doc):
    heading(doc, "5. Frontend Buttons and Actions with APIs", 1)
    paragraph(doc, "This section explains what each important ProjectHub frontend button/action does and which backend API it uses.", after=6)
    for screen, actions in BUTTON_ACTIONS.items():
        heading(doc, screen, 2)
        table(doc, ["Button / Action", "API Used", "Meaning"], actions, widths=[1.8, 2.2, 2.5])


def add_conclusion(doc):
    heading(doc, "6. Simple End-to-End Explanation", 1)
    for item in [
        "Student logs in and receives JWT token.",
        "Frontend stores token and sends it in Authorization header for protected pages.",
        "Student submits project/certificate; file is uploaded first, then project/certificate record is saved.",
        "Faculty uses assigned project/certificate APIs to approve or reject submissions.",
        "Admin manages users, settings, deadlines, faculty assignments, and all submissions.",
        "Notifications inform users about submission, review, and status changes.",
    ]:
        bullet(doc, item)


def configure_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12


def main():
    entries, _ = build_entries()
    doc = Document()
    configure_doc(doc)
    add_title(doc)
    add_annotations(doc)
    add_jwt(doc)
    add_api_summary(doc, entries)
    add_api_details(doc, entries)
    add_buttons(doc)
    add_conclusion(doc)
    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX}")
    print(f"apis {len(entries)}")
    print(f"button_sections {len(BUTTON_ACTIONS)}")


if __name__ == "__main__":
    main()
